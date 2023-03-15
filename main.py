# modules
from docx import Document
from docx.opc.coreprops import CoreProperties
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from tkinter import *
from tkinter import messagebox


ws = Tk()
ws.title('Service Certificate Generator')
ws.geometry('400x300')
ws.config(bg='#456')

f = ('sans-serif', 13)
btn_font = ('sans-serif', 10)
bgcolor = '#BF5517'

genvar = StringVar()
genopt = ['Male', 'Female']
genvar.set('Male')

def clear_inputs():
    eid.delete(0, 'end')
    ename.delete(0, 'end')
    desig.delete(0, 'end')
    jd.delete(0, 'end')
    ed.delete(0, 'end')
    


def generate():

     # data variables
    logo = 'files/tsinfo_logo.jpg'
    output = 'output/Experience_letter.docx'
    sign = 'files/signature.png'
    ceo_sig_text = '''Adarshnath Singh \nDirector'''
    heading = 'Service Certificate'
    emp_id = eid.get() 
    emp_name = ename.get() 
    designation = desig.get()
    joining_date = jd.get()
    end_date = ed.get()

    comp_detail = '''
    TSInfo Technologies (OPC) Pvt Ltd
    Flat G-115, SJR Residency, Devarabeesanahalli, Bellandur, Bangalore, 560103
    Email: info@tsinfotechnologies.com, Phone: +91-9916854253
    '''
    # gender specification 

    gen1 = 'He' # she
    gen2 = 'his' # her
    gen3 = 'him' # her

    if genvar.get() == 'Male':
        gen1 = 'He'
        gen2 = 'his'
        gen3 = 'him' 
    elif genvar.get() == 'Female':
        gen1 = 'She'
        gen2 = 'her' 
        gen3 = 'her'
    else:
        messagebox.showerror('Error', 'Incorrect gender Selection!')

    # experience certificate template
    body_text = f'''
This is to certify that {emp_name} has worked with TSInfo Technologies (OPC) Pvt Ltd from {joining_date} to {end_date}, and was designated as {designation} at the time of {gen2} leaving the organization.

{gen1} is hardworking and a good team player.

We wish {gen3} all the success in {gen2} future endeavor.

    '''

    # create object(s)
    doc =  Document()
    sections = doc.sections


   # declare margin
    for section in sections:
        section.top_margin = Inches(0.04)
        section.bottom_margin = Inches(0.19)
        section.left_margin = Inches(0.93)
        section.right_margin = Inches(0.89)

    section = doc.sections[0]


    # logo image placement
    logo = doc.add_picture(logo, width=Inches(2.52), height=Inches(0.81))
    logo_placement = doc.paragraphs[-1] 
    logo_placement.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # line space
    for _ in range(1):
        linespace_style = doc.styles['Body Text']
        linespace = doc.add_paragraph(style=linespace_style).add_run(' ')
        linespace_style.font.size = Pt(10)


    # employee Id
    empid_style = doc.styles['Normal']
    empid = doc.add_paragraph(style=empid_style).add_run(f'{emp_id}')
    empid.font.bold = True

    # line space
    for _ in range(1):
        linespace_style = doc.styles['Body Text']
        linespace = doc.add_paragraph(style=linespace_style).add_run()
        linespace.font.size = 10

    # Header 
    heading_style = doc.styles['Body Text']
    head = doc.add_paragraph(style=heading_style).add_run(f'{heading}')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    head.font.size = Pt(20)
    head.font.bold = True 


    # body text 
    body_style = doc.styles['Body Text']
    body = doc.add_paragraph(style=body_style).add_run(f'{body_text}')
    body.font.size = Pt(14)
    body.font.name = 'Times New Roman'

    #line space
    for _ in range(2):
        linespace_style = doc.styles['Body Text']
        linespace = doc.add_paragraph(style=linespace_style).add_run()
        linespace.font.size = 10
  


    # signature image & text
    ceo_sign = doc.styles['Body Text']
    doc.add_picture(sign, width=Inches(1.57), height=Inches(0.43))
    doc.add_paragraph(style=ceo_sign).add_run(f'{ceo_sig_text}')
    ceo_sign.font.size = Pt(14)
    


    # line space
    for _ in range(4):
        linespace_style = doc.styles['Body Text']
        linespace = doc.add_paragraph(style=linespace_style)
        # linespace.font.size = Pt(10)

    # footer text : company description
    company_text = doc.styles['Normal']
    company_text.paragraph_format.space_before = Pt(12)
    doc.add_paragraph(style=company_text).add_run(f'{comp_detail}')
    center_align = doc.paragraphs[-1] 
    center_align.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(output)

# frames
frame = Frame(ws, padx=20, pady=20, bg=bgcolor)
frame.pack(expand=True, fill=BOTH)



# label widgets
Label(
    frame, 
    text="Employee ID",
    font=f,
    bg=bgcolor
).grid(row=0, column=0, sticky='w')

Label(
    frame,
    text="Employee Name",
    font=f,
    bg=bgcolor
).grid(row=1, column=0, sticky='w')

Label(
    frame,
    text="Designation",
    font=f,
    bg=bgcolor
).grid(row=2, column=0, sticky='w')

Label(
    frame,
    text="Joining Date",
    font=f,
    bg=bgcolor
).grid(row=3, column=0, sticky='w')

Label(
    frame,
    text="End Date",
    font=f,
    bg=bgcolor
).grid(row=4, column=0, sticky='w')

Label(
    frame,
    text='Gender',
    font=f,
    bg=bgcolor
).grid(row=5, column=0, sticky='w')



# entry widgets
eid = Entry(frame, width=20, font=f)
eid.grid(row=0, column=1)

ename = Entry(frame, width=20, font=f)
ename.grid(row=1, column=1)

desig = Entry(frame, width=20, font=f)
desig.grid(row=2, column=1)

jd = Entry(frame, width=20, font=f)
jd.grid(row=3, column=1)

ed = Entry(frame, width=20, font=f)
ed.grid(row=4, column=1)

gender = OptionMenu(
    frame, 
    genvar,
    *genopt
)
gender.grid(row=5, column=1, pady=(5,0))
gender.config(width=15, font=f)


btn_frame = Frame(frame, bg=bgcolor)
btn_frame.grid(columnspan=2, pady=(50, 0))

# default inputs for testing
eid.insert(0,'01')
ename.insert(0, 'Olamide Omitigun')
desig.insert(0, 'Python Developer')
jd.insert(0, 'April 11st, 2022 ')
ed.insert(0, 'At work')



submit_btn = Button(
    btn_frame,
    text='Generate Word',
    command=generate,
    font=btn_font,
    padx=10, 
    pady=5
)
submit_btn.pack(side=LEFT, expand=True, padx=(15, 0))

clear_btn = Button(
    btn_frame,
    text='Clear',
    command=clear_inputs,
    font=btn_font,
    padx=10, 
    pady=5,
    width=7
)
clear_btn.pack(side=LEFT, expand=True, padx=15)

exit_btn = Button(
    btn_frame,
    text='Exit',
    command=lambda:ws.destroy(),
    font=btn_font,
    padx=10, 
    pady=5
)
exit_btn.pack(side=LEFT, expand=True)


# mainloop
ws.mainloop()
